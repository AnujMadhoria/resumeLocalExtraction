"""Create local, text-based dummy resumes for extractor testing."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate


ROOT = Path(__file__).resolve().parents[1]
PDF_DIR = ROOT / "output" / "pdf"
DOCX_DIR = ROOT / "output" / "docx"
NAVY = "17324D"
BLUE = "2E74B5"
GRAY = "5F6B76"


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "name": ParagraphStyle("ResumeName", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=22, leading=26, textColor=colors.HexColor(f"#{NAVY}"), spaceAfter=3),
        "role": ParagraphStyle("ResumeRole", parent=base["Normal"], fontName="Helvetica", fontSize=10.5, leading=14, textColor=colors.HexColor(f"#{GRAY}"), spaceAfter=7),
        "contact": ParagraphStyle("ResumeContact", parent=base["Normal"], fontName="Helvetica", fontSize=8.8, leading=12, textColor=colors.HexColor("#333333"), spaceAfter=10),
        "heading": ParagraphStyle("ResumeHeading", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=10.5, leading=14, textColor=colors.HexColor(f"#{BLUE}"), spaceBefore=9, spaceAfter=4),
        "body": ParagraphStyle("ResumeBody", parent=base["Normal"], fontName="Helvetica", fontSize=9.2, leading=13, textColor=colors.HexColor("#202020"), spaceAfter=3, alignment=TA_LEFT),
        "job": ParagraphStyle("ResumeJob", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=9.5, leading=13, textColor=colors.HexColor(f"#{NAVY}"), spaceBefore=2, spaceAfter=1),
    }


def build_pdf(path: Path, data: dict) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(path), pagesize=LETTER, rightMargin=0.72 * inch, leftMargin=0.72 * inch,
        topMargin=0.58 * inch, bottomMargin=0.58 * inch,
        title=f"Dummy Resume - {data['name']}", author="Resume Extractor Test Data",
    )
    story = [
        Paragraph(data["name"], styles["name"]),
        Paragraph(data["role"], styles["role"]),
        Paragraph(data["contact"], styles["contact"]),
        HRFlowable(width="100%", thickness=0.8, color=colors.HexColor(f"#{BLUE}"), spaceAfter=4),
    ]
    for section, items in data["sections"]:
        content = [Paragraph(section, styles["heading"])]
        for item in items:
            style = styles["job"] if item.startswith("ROLE: ") else styles["body"]
            content.append(Paragraph(item.removeprefix("ROLE: "), style))
        # Each resume is intentionally one page; adding flowables separately
        # prevents a KeepTogether split from orphaning a section heading.
        story.extend(content)
    doc.build(story)


def set_font(run, *, name: str = "Aptos", size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def add_docx_section(doc: Document, heading: str, entries: list[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(heading)
    set_font(run, size=10.5, color=BLUE, bold=True)

    for label, text in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.08
        if label:
            run = p.add_run(label)
            set_font(run, size=9.4, color=NAVY, bold=True)
        run = p.add_run(text)
        set_font(run, size=9.4, color="202020")


def build_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(9.4)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("DUMMY RESUME | RESUME EXTRACTOR TEST DATA")
    set_font(run, size=7.5, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Fictional candidate profile - created for local testing")
    set_font(run, size=7.5, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("NORA WILLIAMS")
    set_font(run, size=22, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run("Product Designer | Mobile and SaaS Experiences")
    set_font(run, size=10.5, color=GRAY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("nora.williams@example.test | +1 (415) 555-0186 | linkedin.com/in/nora-williams | github.com/norawilliams")
    set_font(run, size=8.8, color="303030")

    add_docx_section(doc, "SUMMARY", [
        ("", "Product designer with 5 years of experience shaping accessible web and mobile workflows for B2B teams."),
    ])
    add_docx_section(doc, "SKILLS", [
        ("", "Figma | User Research | Wireframing | Prototyping | Design Systems | HTML | CSS | JavaScript | Agile | Jira"),
    ])
    add_docx_section(doc, "WORK EXPERIENCE", [
        ("Senior Product Designer | Brightside Systems | ", "Apr 2023 - Present"),
        ("", "Led research and end-to-end design for a workflow platform used by 40,000 operations users."),
        ("Product Designer | Harbor Labs | ", "Jul 2020 - Mar 2023"),
        ("", "Created a reusable component library and improved task completion in the onboarding flow."),
    ])
    add_docx_section(doc, "EDUCATION", [
        ("Bachelor of Design in Interaction Design | ", "Northshore Institute of Design | 2020"),
    ])
    add_docx_section(doc, "CERTIFICATIONS", [
        ("", "Google UX Design Certificate | 2021"),
    ])
    doc.core_properties.title = "Dummy Resume - Nora Williams"
    doc.core_properties.author = "Resume Extractor Test Data"
    doc.save(path)


def main() -> None:
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    DOCX_DIR.mkdir(parents=True, exist_ok=True)

    build_pdf(PDF_DIR / "maya_sharma_data_scientist.pdf", {
        "name": "MAYA R. SHARMA",
        "role": "Data Scientist | Machine Learning and Analytics",
        "contact": "maya.sharma@example.test | +91 98765 43210 | linkedin.com/in/maya-r-sharma | github.com/mayasharma",
        "sections": [
            ("SUMMARY", ["Data scientist with 4 years of experience building forecasting and classification solutions for retail operations."]),
            ("SKILLS", ["Python | SQL | Machine Learning | Deep Learning | Pandas | NumPy | scikit-learn | TensorFlow | Power BI | AWS | Docker | Git"]),
            ("WORK EXPERIENCE", [
                ("ROLE: Data Scientist | Meridian Retail Analytics | Jun 2023 - Present"),
                ("Built demand forecasting models that improved inventory planning accuracy across regional stores."),
                ("ROLE: Data Analyst | QuantBay Solutions | Jul 2021 - May 2023"),
                ("Developed SQL dashboards and automated data-quality checks for client reporting."),
            ]),
            ("EDUCATION", ["M.Tech in Data Science | Institute of Technology, Pune | 2021", "B.Tech in Computer Science | Lakeside University | 2019"]),
            ("CERTIFICATIONS", ["AWS Certified Cloud Practitioner | 2024", "Google Data Analytics Professional Certificate | 2022"]),
        ],
    })

    build_pdf(PDF_DIR / "arjun_mehta_backend_engineer.pdf", {
        "name": "ARJUN MEHTA",
        "role": "Backend Engineer | Distributed Systems",
        "contact": "arjun.mehta@example.test | +91 99887 66554 | https://www.linkedin.com/in/arjunmehta | https://github.com/arjun-mehta",
        "sections": [
            ("PROFILE", ["Backend engineer focused on secure APIs, event-driven services, and cloud infrastructure for financial products."]),
            ("TECHNICAL SKILLS", ["Java | Python | SQL | Spring Boot | REST APIs | PostgreSQL | Redis | Kafka | AWS | Kubernetes | Terraform | Jenkins | CI/CD | Git"]),
            ("PROFESSIONAL EXPERIENCE", [
                ("ROLE: Senior Backend Engineer | OrbitPay Technologies | Mar 2022 - Present"),
                ("Designed high-availability payment services and reduced API latency through cache and query optimization."),
                ("ROLE: Software Engineer | Northwind Digital | Aug 2019 - Feb 2022"),
                ("Implemented Java services, database migrations, and automated deployment pipelines."),
            ]),
            ("EDUCATION", ["B.E. in Information Technology | Greenfield College of Engineering | 2019"]),
            ("PROJECTS", ["Ledger Stream: Kafka-based audit event service with schema validation and replay tooling."]),
        ],
    })

    build_docx(DOCX_DIR / "nora_williams_product_designer.docx")
    print("Created sample resumes:")
    for path in sorted(PDF_DIR.glob("*.pdf")):
        print(path)
    for path in sorted(DOCX_DIR.glob("*.docx")):
        print(path)


if __name__ == "__main__":
    main()
